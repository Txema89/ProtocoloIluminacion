from docx import Document
from docx.shared import Pt
from docx.shared import * 
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL


print("Protocolo de iluminacion by Txema/Zapi")

empresa = 'Corredores Viales SA' #input("Ingrese nombre de la empresa ")
cuit = '35030043' #input("Ingrese CUIT de la empresa ")
dire = 'Pte. Saenz Pena 777' #input("Ingrese direccion de la empresa ")
localidad = 'CABA' #input("Ingrese localidad de la empresa ")
provincia = 'Bs As' #input("Ingrese provincia de la empresa ")
cp = '2000' #input("Ingrese codigo postal ")
instrumento = 'marca TES, modelo 1330, serie 92409619'#input("Ingrese Marca y Modelo del intrumento de Medicion ")
calibrado = '05/08/2023'#input("Ingrese Fecha de Calibracion ")
puestos = 2 #int(input("Ingrese la cantidad de puestos de trabajo a evaluar "))
data_puesto = {} # CREA DICCIONARIO SIN DATOS
tabla_puestos = [] #  CREA LISTA DE DICCIONARIOS EN BASE A DATA_SOURCE


def med_puesto (nombre_puesto, renglon):                          # DECLARA ARRAY SOLICITA MEDIDAS DEL PUESTO X,Y Y GENERA UNA MATRIZ
    """Solicita nombres de los puestos y caracteristicas necesarias para realizar los calculos, declara array, genera matriz y llama a la funcion claculos() para guardar en un diccionario toda la informacion de la medicion del puesto 


    Args:
        nombre_puesto (STR): Nombre del puesto como clave
        renglon (INT): "renglon" cantidad de iteraciones

    Returns:
        DICCT: medicion_puesto con toda la data con clave : valor
    """
    sector = 'of. Central' #str(input (f'Ingrese Sector de {nombre_puesto}: -> '))
    x = 1 #int(input("Ingrese medida de FILAS 'X': "))  
    y = 2 #int(input("Ingrese medida de COLUMNAS 'Y': "))
    a = '3.50m'  #str(input("Ingrese la altura de la Iluminaria: "))
    iluminacion = 'General' #str(input (f'Ingrese Ilumunacion de {nombre_puesto}: General / Localizada / Mixta -> '))
    tiluminacion = 'Artificial' #str(input (f'Ingrese el Tipo de Ilumunacion de {nombre_puesto}: Natural / Artificial / Mixta -> '))
    tluminica = 'Mixta' #str(input (f'Ingrese el Tipo de Fuente Lumínica de {nombre_puesto}: Incandescente / Descarga / Mixta -> '))
    vrequerido = 666 #int(input (f'Ingrese el Valor Requerido de {nombre_puesto} -> '))
    

    lux = [
            [0 for j in range(y)]
            for i in range(x)
        ] # Declara Array de dimenciones X; magia!!

    for i in range (0,x):
        for j in range (0,y):
                while True:
                    try:    
                        valor = float (input ("Ingrese el valor de X: " + str(i+1) + " ,Y: " + str(j+1) + "--> ")) # CORREJIR ERROR; PARA NO ROMPER ITERACION SINO SE INGRESA VALOR NUMERICO
                        lux [i][j] = valor
                        break
                    except ValueError:
                        print("Valor incorrecto!")
    
    minimo, media = calculos(lux, x, y)
    if minimo < (media / 2):
        masmenos = '<'
    elif minimo > (media / 2):
        masmenos = '>'
    else:
        masmenos = '='            
    mediciones_puesto = {
        'muestreo' : str(renglon + 1),
        'hora' : '5pm' ,
        'sector' : str(sector),
        'puesto' : str(nombre_puesto),
        'tiluminacion' : str(tiluminacion),
        'tluminica' : str(tluminica),
        'iluminacion' : str(iluminacion),             
        'minimo' : str(minimo),
        'masmenos' : str(masmenos),
        'media/2' : str(media/2),
        'media' : str(media),
        'vrequerido' : str(vrequerido),
        #'lux' : lux,
        #'x' : x,
        #'y' : y,
    }
    tabla_puestos.append(mediciones_puesto)
    return lux, x, y, mediciones_puesto


def calculos(lux, x, y):                    # CALCULA EN LA MATRIZ LA MADIA Y COMPARA CON EL VALOR MINIMO (UNIFORMIDAD DE ILUMINACION)
    """Esta funcion es llamada en med_puesto para realizar los calculos y comparaciones; recibe los valores de x, y, lux para realizar calculos y obtener valores como la media aritmetica 'media' y el valor mas chico 'minimo' requeridos en el informe.

    Args:
        lux (LIST): lista de listas, matriz
        x (INT): cantidad de filas 
        y (INT): cantidad de columnas

    Returns:
        INT: minimo
        INT: media
        DICT: mediciones_puesto
    """

    for r in lux:
        print(r) 

    sumatoria = 0 
    minimo = 0

    for i in range (0, x):
        for j in range (0, y):
            valor = lux[i][j]
            sumatoria += valor
            if i == 0 and j == 0:
                minimo = valor 
            elif valor < minimo:
                minimo = valor
    media = sumatoria / (x*y)
    print ("La media es " + str(round(media,2)))
    print ("El minimo es " + str(minimo))
    print (f"Cantidad de Mediciones {x*y}")
    if minimo < (media / 2):
        print ("El valor minimo es mas chico que la media! La uniformidad no es adecuada.")
    return minimo, media
    

def change_orientation(document):           # CAMBIA ORIENTACION DE PAGINA EN CREAR DOC
    """Cambia la orietacion de la hoja de Word. XD

    Args:
        document (_type_??): _description_

    Returns:
        _type_??: Nueva seccion en el document
    """
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


def crear_doc(empresa:str, data_puesto:dict):        # CREA DOCUMENTO WORD
    """Crea documento Word

    Args:
        empresa (str): Nombre de la empresa
        data_puesto (dicct): diccionario con clave nombre del puesto y valor lux(lista de listas; matriz)
        tabla_puestos (dicct): diccionario con toda la info para la tabla descriptiva en el documento
    """

    document = Document()

    # Working with Sections (That is de only Section for all document... for now...)
    section = document.sections[0]
    # Selecting the header
    header = section.header
    # Selecting the paragraph already present in the header section
    header_para = header.paragraphs[0]

    # Adding the centred zoned header
    header_para.text = f"\t{empresa} - CUIT {cuit} - Ciudad: {localidad} - Dirección: {dire} "

    footer = section.footer
    # Calling the paragraph already present in the footer section
    footer_para = footer.paragraphs[0]
    # Adding text in the footer
    footer_para.text = "\t\t " #TABULACION PARA LA FIRMA
    run = footer_para.add_run()
    run.add_picture('FIRMA digital.png',width=Inches(1.5), height=Inches(1.5))

    # DEFINICION DE UN ESTILO 
    styles = document.styles
    charstyle = styles.add_style('Calibri_18', WD_STYLE_TYPE.CHARACTER) #DEFINE EL NOMBRE DEL ESTICO COMO "CALIBRI_18"
    obj_font = charstyle.font
    obj_font.size = Pt(18)
    obj_font.name = 'Calibri'

    # CONTENIDO DEL INFORME
    # CARATULA

    title = document.add_heading('Protocolo de Iluminación Resolución SRT 84/2012', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # INFORMACION DE LA EMPRESA

    document.add_paragraph('\n\n\n\n') #CENTRANDO EL TEXTO XD
    document.add_heading('INFORMACION DE LA EMPRESA\n', 2)

    paragraph = document.add_paragraph()
    paragraph.add_run('\tNombre de la empresa: '+ '\t\t' + empresa + '\n', style = 'Calibri_18')
    paragraph.add_run('\tNumero de CUIT/CUIL: '+ '\t\t' +cuit + '\n', style = 'Calibri_18')
    paragraph.add_run('\tDirección: '+ '\t\t\t\t' +dire + '\n', style = 'Calibri_18')
    paragraph.add_run('\tLocalidad: '+ '\t\t\t\t' +localidad + '\n', style = 'Calibri_18')
    paragraph.add_run('\tCódigo Postal: '+ '\t\t\t' +cp + '\n\n\n\n', style = 'Calibri_18')
    
    document.add_heading('PROFESIONAL INTERVINIENTE\n', 2)
    paragraph = document.add_paragraph()
    paragraph.add_run('\tJose M Perez Garcia', style = 'Calibri_18')

    document.add_page_break()

    # HOJA DESCRIPTIVA

    title = document.add_heading('PORTADA DESCRIPTIVA DE PROTOCOLOS “NIVELES DE ILUMINACIÓN”', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('PROTOCOLO PARA MEDICION DE ILUMINACION DE AMBIENTE LABORAL', level=2)
    document.add_paragraph(f'''
        Razón social: {empresa}
        Dirección: {dire}
        Localidad: {localidad}
        Provincia: {provincia}
        CP: {cp}
        CUIT: {cuit}
    ''')
    
    document.add_heading('TURNOS DE TRABAJO:', level=2)
    document.add_paragraph(f'''
        Administración:
        Personal Operativo:
    ''')

    document.add_heading('DATOS DE LA MEDICION', level=2)
    document.add_paragraph(f'''
    Marca, modelo y número de serie del instrumento utilizado: {instrumento}
                           
    Fecha de Calibración del Instrumental utilizado en la medición: {calibrado}
                           
    Metodología Utilizada en la Medición: Se utilizó el método de la cuadrícula propuesto en la Resolución SRT 84/2012
    
    Fecha de la Medición:
    Hora de Inicio:
    Hora de Finalización:
    Condiciones atmosféricas:
    ''')
    
    document.add_heading('DOCUMENTACION QUE SE ADJUNTA A LA MEDICION', level=2)
    document.add_paragraph(f'''
    Certificado de Calibración: 
    Croquis del establecimiento:
    Observaciones:''')    

    #DATOS DE LA MEDICION EN TABLA Y NUEVA SECCION APAISADA 
    # Configurar la orientación de la página en apaisada
    change_orientation(document)

    document.add_heading('DATOS DE LA MEDICION', level=0)
    

    table = document.add_table(rows = 1, cols = 12, style = 'Table Grid')
    

    encabezados = ['Punto de Muestreo', 'Hora', 'Sector', 'Sección /Puesto /Puesto Tipo', 
           'Tipo de Iluminación: Natural /Artificial /Mixta', 'Tipo de Fuente Lumínica: Incandescente /Descarga /Mixta', 
           'Iluminación: General /Localizada /Mixta', 'V.Unif E.min', '', 'V.Unif E.med/2', 'Valor Medido (Lux)', 
           'Valor requerido legalmente Según Anexo IV Dec. 351/79']

    # Agregar los encabezados a la primera fila de la tabla
    for i, encabezados in enumerate(encabezados):
        table.cell(0, i).text = encabezados
        # Alineación vertical centrada para los encabezados
        cell = table.cell(0, i)
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(9)

    table = document.add_table(rows = puestos, cols = 12, style = 'Table Grid')


    fila = 0  

    for i in tabla_puestos:
        columna = 0  

        for key, value in i.items():
            cell = table.cell(fila, columna)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(12)
            
            columna += 1  

        fila += 1  
    
    document.add_page_break()
   
    # CONCLICIONES Y MEJORAS

    document.add_heading('Análisis de los Datos y Mejoras a Realizar', level=0)
    
    table = document.add_table(rows = 2, cols = 2, style = 'Table Grid')
    table.cell(0, 0).text = 'Concluciones:'
    table.cell(0, 1).text = 'Recomendaciones para adecuar el nivel de iluminación a la legislación vigente:'
    table.cell(1, 1).text = 'Se deberá implementar un programa de mantenimiento periódico preventivo y de limpieza de luminarias, a fin de detectar y corregir anormalidades.'
    
    # MATRICES

    document.add_page_break()
    change_orientation(document)

    for n_puesto in data_puesto.keys(): # HACIENDO GETs 
        lux = data_puesto[n_puesto][0]
        x = data_puesto[n_puesto][1]
        y = data_puesto[n_puesto][2]
        document.add_paragraph({n_puesto}, style='Intense Quote')
        table = document.add_table(rows = 0, cols = y, style = 'Light Shading') # INICIA LAS FILAS EN 0 
        for row in range(x):
            row_cells = table.add_row().cells
            for colum in range(y):
                row_cells[colum].text = str(lux[row][colum])
        sumatoria = 0 
        minimo = 0

        for i in range (0, x):
            for j in range (0, y):
                valor = lux[i][j]
                sumatoria += valor
                if i == 0 and j == 0:
                    minimo = valor
                elif valor < minimo:
                    minimo = valor
        media = sumatoria / (x*y)
        document.add_paragraph("La media es " + str (round(media,2)))
        document.add_paragraph("El minimo es " + str(minimo))
        document.add_paragraph(f"Cantidad de Mediciones {x*y}")
        if minimo < (media / 2):
            document.add_paragraph (f"El valor minimo es mas chico que la media/2 ({media/2}) - La uniformidad NO es adecuada - ")
    
    document.save('Protocolo de Iluminacion para la empresa ' + empresa + '.docx')


for puesto in range (puestos):          # BUCLE QUE SOLICITA CANTIDAD DE PUESTOS Y REPITE LAS MEDICIONES Y CALCULOS DE MATRICES
    n_puesto = input("Ingrese nombre del puesto de trabajo -> ")
    medicion_puesto = med_puesto(n_puesto, puesto) 
    data_puesto[n_puesto] = medicion_puesto #DICCIONARIO CON CLAVE 'NOMBRE DEL PUESTO' QUE CONTIENE UN DICCIONARIO.
    
crear_doc (empresa, data_puesto)